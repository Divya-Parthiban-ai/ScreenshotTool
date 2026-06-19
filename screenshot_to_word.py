import pyautogui
from docx import Document
from docx.shared import Inches
import datetime
import keyboard
import os
import tkinter as tk
from tkinter import simpledialog

# === BASE FOLDER ===
BASE_FOLDER = r"C:\Users\qev5128\TestEvidence"

# === GET USE CASE NAME ===
use_case = input("Enter Use Case Name: ").strip()

# === CREATE FOLDERS ===
USECASE_FOLDER = os.path.join(BASE_FOLDER, use_case)
IMG_FOLDER = os.path.join(USECASE_FOLDER, "Screenshots")
DOC_PATH = os.path.join(USECASE_FOLDER, "TestReport.docx")

os.makedirs(IMG_FOLDER, exist_ok=True)

# === LOAD OR CREATE WORD DOC ===
if os.path.exists(DOC_PATH):
    doc = Document(DOC_PATH)
else:
    doc = Document()
    doc.add_heading(f'Test Evidence Report - {use_case}', 0)

step_counter = 1

# === CAPTURE FUNCTION ===
def capture():
    global step_counter

    # ✅ TIMESTAMP + IMAGE PATH
    timestamp = datetime.datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    img_path = os.path.join(IMG_FOLDER, f"shot_{timestamp}.png")

    # ✅ STEP 1: TAKE SCREENSHOT FIRST (clean)
    screenshot = pyautogui.screenshot()
    screenshot.save(img_path)

    # ✅ STEP 2: POPUPS AFTER SCREENSHOT
    root = tk.Tk()
    root.withdraw()

    step_desc = simpledialog.askstring("Step Description", "Enter step description:")
    result = simpledialog.askstring("Result", "Enter result (Pass/Fail):")

    root.destroy()

    # ✅ DEFAULT HANDLING
    if not step_desc:
        step_desc = "No Description"

    if result:
        result = result.strip().upper()
    else:
        result = "UNKNOWN"

    if result not in ["PASS", "FAIL"]:
        result = "UNKNOWN"

    # ✅ WRITE TO WORD
    doc.add_paragraph(f"Step {step_counter}: {step_desc}")
    doc.add_paragraph(f"Time: {timestamp} | Result: {result}")
    doc.add_picture(img_path, width=Inches(5))
    doc.add_paragraph("---------------------------------------------------")

    # ✅ SAFE SAVE
    try:
        doc.save(DOC_PATH)
    except PermissionError:
        print("⚠️ Close the Word file and try again.")
        return

    print(f"✅ Step {step_counter} captured - {result}")

    step_counter += 1


# === HOTKEY ===
keyboard.add_hotkey('ctrl+alt+s', capture)

print("\n✅ Tool Ready")
print(f"📂 Saving to: {USECASE_FOLDER}")
print("👉 Press Ctrl + Alt + S to capture screenshots")

keyboard.wait()